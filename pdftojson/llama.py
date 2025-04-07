import json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import re
from bs4 import BeautifulSoup
from docx.shared import Pt, RGBColor
from html import unescape

# Function to clean HTML content
def clean_html_content(text):
    if not isinstance(text, str):
        return str(text)
    # Remove HTML tags and decode HTML entities
    text = re.sub(r'<[^>]+>', '', text)
    text = unescape(text)
    return text

# Function to add a heading to Word document
def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

# Function to add a paragraph to Word document
def add_paragraph(doc, text):
    para = doc.add_paragraph(text)
    return para

def apply_html_style_to_run(run, tag):
    if tag.name == 'strong' or tag.name == 'b':
        run.bold = True
    if tag.name == 'em' or tag.name == 'i':
        run.italic = True
    if tag.name == 'u':
        run.underline = True
    if tag.get('style'):
        style = tag['style']
        if 'text-decoration: underline' in style:
            run.underline = True

def add_html_paragraph(doc, html_content):
    if not isinstance(html_content, str):
        return doc.add_paragraph(str(html_content))
    
    paragraph = doc.add_paragraph()
    if not html_content.strip():
        return paragraph
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.children:
        if isinstance(element, str):
            run = paragraph.add_run(element)
        else:
            run = paragraph.add_run(element.get_text())
            apply_html_style_to_run(run, element)
    
    return paragraph

# Function to process JSON and export to Word
def export_to_word(data, output_file):
    doc = Document()
    
    # Title
    add_heading(doc, "Preschool2024 Curriculum", 0)

    # Iterate through folders
    for folder in data["folders"]:
        add_heading(doc, folder["name"], 1)
        
        # Add folder details
        add_paragraph(doc, f"ID: {folder['id']}")
        if folder["parent_id"]:
            add_paragraph(doc, f"Parent ID: {folder['parent_id']}")
        add_paragraph(doc, f"Program ID: {folder['program_id']}")
        
        # Process resources if they exist
        if "resources" in folder:
            for resource in folder["resources"]:
                add_heading(doc, resource["title"], 2)
                add_paragraph(doc, f"ID: {resource['id']}")
                add_paragraph(doc, f"Type: {resource['type']}")
                add_paragraph(doc, f"Focus Area: {resource['focus_area']}")
                add_html_paragraph(doc, f"Short Description: {resource['short_description']}")
                
                # Process activities
                if "activities" in resource:
                    add_heading(doc, "Activities", 3)
                    for activity in resource["activities"]:
                        add_heading(doc, activity["title"], 4)
                        add_paragraph(doc, f"ID: {activity['id']}")
                        add_paragraph(doc, f"Type: {activity['type']}")

                        # Process content blocks
                        if "content_blocks" in activity:
                            add_heading(doc, "Content Blocks", 5)
                            for block in activity["content_blocks"]:
                                block_title = block.get("title", block.get("type", ""))
                                add_heading(doc, block_title, 6)
                                add_paragraph(doc, f"Type: {block['type']}")
                                if "content" in block:
                                    if isinstance(block["content"], dict):
                                        if "text" in block["content"]:
                                            add_html_paragraph(doc, block["content"]["text"])
                                        if "title" in block["content"]:
                                            add_html_paragraph(doc, f"Title: {block['content']['title']}")
                                        if "columns" in block["content"]:
                                            for col in block["content"]["columns"]:
                                                add_html_paragraph(doc, f"{col.get('title', '')}: {col['text']}")
                                    elif isinstance(block["content"], list):
                                        for item in block["content"]:
                                            if isinstance(item, dict):
                                                if "text" in item:
                                                    add_html_paragraph(doc, item["text"])
                                                if "title" in item:
                                                    add_html_paragraph(doc, f"Title: {item['title']}")
                                                if "columns" in item:
                                                    for col in item["columns"]:
                                                        add_html_paragraph(doc, f"{col.get('title', '')}: {col['text']}")

                        # Process standards coding schemes and statements
                        if "standards" in activity:
                            add_heading(doc, "Standards", 7)
                            for standard in activity["standards"]:
                                add_paragraph(doc, f"{standard['code']}: {standard['statement']}")

    # Save the Word document
    doc.save(output_file)
    print(f"Word document saved as {output_file}")

# Function to process JSON and export to PDF
def export_to_pdf(data, output_file):
    doc = SimpleDocTemplate(output_file, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create custom styles for HTML content
    styles.add(ParagraphStyle(
        name='HTMLStyle',
        parent=styles['Normal'],
        alignment=TA_LEFT,
        spaceAfter=10
    ))
    
    story = []

    # Title
    story.append(Paragraph("Preschool 2024 Curriculum", styles["Title"]))
    story.append(Spacer(1, 12))

    # Iterate through folders
    for folder in data["folders"]:
        story.append(Paragraph(folder["name"], styles["Heading1"]))
        story.append(Paragraph(f"ID: {folder['id']}", styles["Normal"]))
        if folder["parent_id"]:
            story.append(Paragraph(f"Parent ID: {folder['parent_id']}", styles["Normal"]))
        story.append(Paragraph(f"Program ID: {folder['program_id']}", styles["Normal"]))
        story.append(Spacer(1, 12))

        # Process resources if they exist
        if "resources" in folder:
            story.append(Paragraph("Resources", styles["Heading2"]))
            for resource in folder["resources"]:
                story.append(Paragraph(resource["title"], styles["Heading3"]))
                story.append(Paragraph(f"ID: {resource['id']}", styles["Normal"]))
                story.append(Paragraph(f"Type: {resource['type']}", styles["Normal"]))
                story.append(Paragraph(f"Focus Area: {resource['focus_area']}", styles["Normal"]))
                story.append(Paragraph(clean_html_content(resource["short_description"]), styles["Normal"]))

                # Content Blocks
                if "content_blocks" in resource:
                    story.append(Paragraph("Content Blocks", styles["Heading4"]))
                    for block in resource["content_blocks"]:
                        story.append(Paragraph(f"Type: {block['type']}", styles["Normal"]))
                        if "content" in block and isinstance(block["content"], dict):
                            if "text" in block["content"]:
                                story.append(Paragraph(clean_html_content(block["content"]["text"]), styles["Normal"]))
                            if "title" in block["content"]:
                                story.append(Paragraph(clean_html_content(block["content"]["title"]), styles["Normal"]))
                            if "columns" in block["content"]:
                                for col in block["content"]["columns"]:
                                    story.append(Paragraph(clean_html_content(f"{col['title']}: {col['text']}"), styles["Normal"]))
                story.append(Spacer(1, 12))
        story.append(PageBreak())

    # Build the PDF
    doc.build(story)
    print(f"PDF saved as {output_file}")

# Main function to load JSON and export
def main():
    # Load your JSON data (replace with your file path or string)
    with open("craft_data.json", "r") as f:
        data = json.load(f)
    
    # Export to Word
    export_to_word(data, "preschool_curriculum_llama.docx")
    
    # Export to PDF
    export_to_pdf(data, "preschool_curriculum_llama.pdf")

if __name__ == "__main__":
    main()