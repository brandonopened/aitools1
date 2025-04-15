import json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.colors import orange
import re
from bs4 import BeautifulSoup
from docx.shared import Pt, RGBColor
from html import unescape
import sys

# Function to clean HTML content
def clean_html_content(text):
    if not isinstance(text, str):
        return str(text)
    # Remove HTML tags and decode HTML entities
    text = re.sub(r'<[^>]+>', '', text)
    text = unescape(text)
    return text

# Function to add a heading to Word document
def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

# Function to add a paragraph to Word document
def add_paragraph(doc, text):
    para = doc.add_paragraph(text)
    return para

def apply_html_style_to_run(run, tag):
    if tag.name == 'strong' or tag.name == 'b':
        run.bold = True
    if tag.name == 'em' or tag.name == 'i':
        run.italic = True
    if tag.name == 'u':
        run.underline = True
    if tag.get('style'):
        style = tag['style']
        if 'text-decoration: underline' in style:
            run.underline = True

def add_html_paragraph(doc, html_content):
    if not isinstance(html_content, str):
        return doc.add_paragraph(str(html_content))
    
    paragraph = doc.add_paragraph()
    if not html_content.strip():
        return paragraph
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.children:
        if isinstance(element, str):
            run = paragraph.add_run(element)
        else:
            run = paragraph.add_run(element.get_text())
            apply_html_style_to_run(run, element)
    
    return paragraph

# Helper to remove problematic span styles for ReportLab
def preprocess_html_for_reportlab(html_text):
    if not isinstance(html_text, str):
        return str(html_text)
    # Remove <span style=...> tags specifically, as they cause findSpanStyle error
    # This regex is basic and might need refinement depending on exact HTML patterns
    processed_text = re.sub(r'<span\s+style=.*?>(.*?)</span>', r'\1', html_text, flags=re.IGNORECASE | re.DOTALL)
    # Optionally remove other problematic tags if encountered
    # processed_text = re.sub(r'<para>(.*?)</para>', r'\1', processed_text, flags=re.IGNORECASE | re.DOTALL) # Remove <para> if needed
    return processed_text

# --- Preprocessing --- 
def preprocess_standards(data):
    """ Creates a mapping from resource_id to a list of standard strings. """
    print("--- Starting Standard Preprocessing ---") # DEBUG
    standard_map = {}
    
    # --- Data Gathering: Look in multiple potential locations --- 
    all_potential_standards = []
    
    # 1. Look for top-level 'standards' array
    top_level_standards = data.get("standards", [])
    if top_level_standards:
        all_potential_standards.extend(top_level_standards)
        print(f"Found {len(top_level_standards)} top-level standards in 'standards' array.") # DEBUG
        
    # 2. Look for top-level 'framework_items' array (based on user JSON)
    top_level_framework_items = data.get("framework_items", [])
    if top_level_framework_items:
        all_potential_standards.extend(top_level_framework_items)
        print(f"Found {len(top_level_framework_items)} top-level items in 'framework_items' array.") # DEBUG
        
    # 3. Look for standards nested within activities
    activity_standards_count = 0 # DEBUG
    for folder in data.get("folders", []):
        for resource in folder.get("resources", []):
             # 4. Also check for framework_items directly under resources?
             resource_framework_items = resource.get("framework_items", [])
             if resource_framework_items:
                  all_potential_standards.extend(resource_framework_items)
                  print(f"Found {len(resource_framework_items)} items in 'framework_items' under resource {resource.get('id')}.") # DEBUG
             
             # Check activities
             for activity in resource.get("activities", []):
                activity_stds = activity.get("standards", [])
                if activity_stds:
                    activity_standards_count += len(activity_stds)
                    all_potential_standards.extend(activity_stds)
                # 5. Also check for framework_items directly under activities?
                activity_framework_items = activity.get("framework_items", [])
                if activity_framework_items:
                    all_potential_standards.extend(activity_framework_items)
                    print(f"Found {len(activity_framework_items)} items in 'framework_items' under activity {activity.get('id')}.") # DEBUG
                    
    if activity_standards_count > 0:
        print(f"Found {activity_standards_count} standards nested within activities.") # DEBUG
    # --- End Data Gathering ---

    # Use dict to ensure uniqueness based on standard/item ID
    unique_standards_dict = {}
    for std in all_potential_standards:
        if isinstance(std, dict) and 'id' in std:
            # Use human_coding_scheme if available and code is missing
            if 'code' not in std and 'human_coding_scheme' in std:
                std['code'] = std['human_coding_scheme']
            # Use full_statement if available and statement is missing
            if 'statement' not in std and 'full_statement' in std:
                std['statement'] = std['full_statement']
                
            unique_standards_dict[std['id']] = std
        else:
            print(f"Warning: Found standard/item without ID or not a dict: {std}") # DEBUG
            
    unique_standards = list(unique_standards_dict.values())
    print(f"Processing {len(unique_standards)} unique standards/items...") # DEBUG

    standards_processed_count = 0 # DEBUG
    links_found_count = 0 # DEBUG
    for standard in unique_standards:
        standard_code = standard.get('code', 'NO_CODE') # DEBUG
        
        # Look for the linkage field
        res_items = standard.get("resources_framework_items")
        
        if res_items is None:
             continue # Skip if no linkage field
             
        # Handle potential variations in structure (dict or list)
        if isinstance(res_items, dict):
            res_items = [res_items]
        elif not isinstance(res_items, list):
            print(f"Warning: Standard {standard_code} has unexpected type for resources_framework_items: {type(res_items)}") # DEBUG
            res_items = []
        
        if not res_items: # Skip if list is empty after handling
            continue
            
        standards_processed_count += 1 # DEBUG (Count standards that HAVE the linkage field)
        found_link_for_this_std = False # DEBUG
        for item in res_items:
             if isinstance(item, dict) and "resource_id" in item:
                resource_id = item["resource_id"]
                
                if standard_code == "P-SE.7.a" and resource_id == 98219:
                    print(f"DEBUG: Found link for P-SE.7.a to resource 98219! Item: {item}")
                    
                std_string = f"{standard.get('code','N/A')}: {standard.get('statement', 'N/A')}"
                if resource_id not in standard_map:
                    standard_map[resource_id] = []
                if std_string not in standard_map[resource_id]: # Avoid duplicates
                     standard_map[resource_id].append(std_string)
                     found_link_for_this_std = True # DEBUG
             else:
                 print(f"Warning: Standard {standard_code} has item in resources_framework_items that is not a dict or lacks 'resource_id': {item}") # DEBUG
                 
        if found_link_for_this_std:
            links_found_count += 1 # DEBUG (Count standards for which we actually added a link)
                     
    print(f"Finished processing standards. {standards_processed_count} standards/items had resources_framework_items field. Found links for {links_found_count} of them.") # DEBUG
    if 98219 in standard_map:
        print("DEBUG: Resource ID 98219 IS present in the final standard_map.")
        print(f"DEBUG: Standards for 98219: {standard_map[98219]}")
    else:
        print("DEBUG: Resource ID 98219 IS NOT present in the final standard_map.")
        
    print("--- Finished Standard Preprocessing ---") # DEBUG
    return standard_map

# --- Export Functions --- 

# Function to process JSON and export to Word
def export_to_word(data, output_file):
    doc = Document()
    standard_map = preprocess_standards(data)
    
    # Title
    add_heading(doc, "Preschool 2024 Curriculum", 0)

    # Iterate through folders
    for folder in data.get("folders", []):
        add_heading(doc, folder.get("name", "N/A"), 1)
        add_paragraph(doc, f"ID: {folder.get('id', 'N/A')}")
        if folder.get("parent_id"):
            add_paragraph(doc, f"Parent ID: {folder['parent_id']}")
        add_paragraph(doc, f"Program ID: {folder.get('program_id', 'N/A')}")
        
        # Process resources if they exist
        if "resources" in folder:
            for resource in folder.get("resources", []):
                resource_id = resource.get("id")
                
                # --- Display Standards First ---
                if resource_id in standard_map:
                    add_heading(doc, "Related Standards", 3) # Heading level 3 for standards
                    for std_string in standard_map[resource_id]:
                        add_paragraph(doc, std_string)
                # --- End Standards ---
                
                add_heading(doc, resource.get("title", "N/A"), 2) # Resource title level 2
                add_paragraph(doc, f"ID: {resource_id if resource_id else 'N/A'}")
                add_paragraph(doc, f"Type: {resource.get('type', 'N/A')}")
                add_paragraph(doc, f"Focus Area: {resource.get('focus_area', 'N/A')}")
                add_html_paragraph(doc, f"Short Description: {resource.get('short_description', '')}")
                
                # Process activities
                if "activities" in resource:
                    add_heading(doc, "Activities", 3) # Activities level 3
                    for activity in resource.get("activities", []):
                        add_heading(doc, activity.get("title", "N/A"), 4) # Activity title level 4
                        add_paragraph(doc, f"ID: {activity.get('id', 'N/A')}")
                        add_paragraph(doc, f"Type: {activity.get('type', 'N/A')}")

                        # Process content blocks
                        if "content_blocks" in activity:
                            add_heading(doc, "Content Blocks", 5) # Content Blocks level 5
                            for block in activity.get("content_blocks", []):
                                block_title = block.get("title", block.get("type", ""))
                                add_heading(doc, block_title, 6) # Block title level 6
                                add_paragraph(doc, f"Type: {block.get('type', 'N/A')}")
                                if "content" in block:
                                    # ... (existing content block processing) ...
                                    content = block["content"]
                                    if isinstance(content, dict):
                                        if "text" in content:
                                            add_html_paragraph(doc, content['text'])
                                        if "title" in content:
                                            # Add title if different from block heading, or skip
                                            if block_title != content['title']:
                                                 add_html_paragraph(doc, f"Content Title: {content['title']}")
                                        if "columns" in content:
                                            for col in content.get("columns", []):
                                                add_html_paragraph(doc, f"{col.get('title', '')}: {col.get('text', '')}")
                                    elif isinstance(content, list):
                                        for item in content:
                                            if isinstance(item, dict):
                                                # Process list item content
                                                if "text" in item:
                                                     add_html_paragraph(doc, item["text"])

                        # Process standards coding schemes and statements (COMMENTED OUT - handled above resource)
                        # if "standards" in activity:
                        #    add_heading(doc, "Standards", 7)
                        #    for standard in activity.get("standards", []):
                        #        add_paragraph(doc, f"{standard.get('code','N/A')}: {standard.get('statement', 'N/A')}")

    # Save the Word document
    doc.save(output_file)
    print(f"Word document saved as {output_file}")

# Function to generate a plain text preview from JSON data
def generate_text_preview(data):
    standard_map = preprocess_standards(data)
    preview_lines = []
    preview_lines.append("Preschool 2024 Curriculum\n" + "="*30)

    # Iterate through folders
    for folder in data.get("folders", []):
        preview_lines.append(f"\n## {folder.get('name', 'N/A')} ##")

        # Process resources if they exist
        if "resources" in folder:
            for resource in folder.get("resources", []):
                resource_id = resource.get("id")
                
                # --- Display Standards First ---
                if resource_id in standard_map:
                    preview_lines.append("\n---- Related Standards ----")
                    for std_string in standard_map[resource_id]:
                        preview_lines.append(f"  - {std_string}")
                # --- End Standards ---
                
                preview_lines.append(f"\n### {resource.get('title', 'N/A')} ###")
                preview_lines.append(f"  Type: {resource.get('type', 'N/A')}")
                preview_lines.append(f"  Focus Area: {resource.get('focus_area', 'N/A')}")
                preview_lines.append(f"\n{clean_html_content(resource.get('short_description', ''))}")
                
                # Process activities (if needed in preview)
                if "activities" in resource:
                    preview_lines.append("\n---- Activities ----")
                    for activity in resource.get("activities", []):
                        preview_lines.append(f"\n#### {activity.get('title', 'N/A')} ####")
                        preview_lines.append(f"  Type: {activity.get('type', 'N/A')}")

                        # Process content blocks (if needed in preview)
                        if "content_blocks" in activity:
                             preview_lines.append("\n  ---- Content Blocks ----")
                             for block in activity.get("content_blocks", []):
                                 block_title = block.get("title", block.get("type", ""))
                                 preview_lines.append(f"\n##### {block_title} #####")
                                 if "content" in block:
                                     # ... (existing preview content block processing) ...
                                     content = block["content"]
                                     if isinstance(content, dict):
                                         if "text" in content:
                                             preview_lines.append(f"\n{clean_html_content(content['text'])}")
                                         # Skip content title if same as block title
                                         if "columns" in content:
                                             preview_lines.append("\n    ---- Columns ----")
                                             for col in content.get("columns", []):
                                                 col_title = clean_html_content(col.get('title', ''))
                                                 col_text = clean_html_content(col.get('text', ''))
                                                 preview_lines.append(f"      {col_title}: {col_text}")
                                     elif isinstance(content, list):
                                         preview_lines.append("\n    ---- List Items ----")
                                         for item in content:
                                             if isinstance(item, dict) and "text" in item:
                                                  preview_lines.append(f"      - {clean_html_content(item['text'])}")
                        
                        # Process standards (COMMENTED OUT - handled above resource)
                        # if "standards" in activity:
                        #     preview_lines.append("\n  ---- Standards ----")
                        #     for standard in activity.get("standards", []):
                        #         preview_lines.append(f"    - {standard.get('code','N/A')}: {standard.get('statement', 'N/A')}")
                                         
        preview_lines.append("\n" + "="*30) # Separator between folders

    return "\n".join(preview_lines)

# Function to process JSON and export to PDF
def export_to_pdf(data, output_file):
    doc = SimpleDocTemplate(output_file, pagesize=letter)
    styles = getSampleStyleSheet()
    standard_map = preprocess_standards(data)
    
    # Set orange color for PDF headings using ReportLab colors
    styles['Heading1'].textColor = orange
    styles['Heading2'].textColor = orange
    styles['Heading3'].textColor = orange
    styles['Heading4'].textColor = orange
    # Add more heading levels if used and style exists
    if 'Heading5' in styles: styles['Heading5'].textColor = orange
    if 'Heading6' in styles: styles['Heading6'].textColor = orange
    if 'Heading7' in styles: styles['Heading7'].textColor = orange
    
    # Optional: Style for standards list
    styles.add(ParagraphStyle(name='StandardStyle', parent=styles['Normal'], leftIndent=18))
    styles.add(ParagraphStyle(name='StandardsHeading', parent=styles['Heading3'], spaceBefore=6, spaceAfter=2))

    story = []

    # Title
    story.append(Paragraph("Preschool 2024 Curriculum", styles["Title"]))
    story.append(Spacer(1, 12))

    # Iterate through folders
    for folder in data.get("folders", []):
        story.append(Paragraph(folder.get("name", "N/A"), styles["Heading1"]))
        story.append(Paragraph(f"ID: {folder.get('id', 'N/A')}", styles["Normal"]))
        if folder.get("parent_id"):
            story.append(Paragraph(f"Parent ID: {folder['parent_id']}", styles["Normal"]))
        story.append(Paragraph(f"Program ID: {folder.get('program_id', 'N/A')}", styles["Normal"]))
        story.append(Spacer(1, 12))

        # Process resources if they exist
        if "resources" in folder:
            for resource in folder.get("resources", []):
                resource_id = resource.get("id")
                
                # --- Display Standards First ---
                if resource_id in standard_map:
                    story.append(Paragraph("Related Standards", styles['StandardsHeading']))
                    for std_string in standard_map[resource_id]:
                        story.append(Paragraph(std_string, styles['StandardStyle']))
                    story.append(Spacer(1, 6))
                # --- End Standards ---
                
                # Resource Title - Heading 2
                story.append(Paragraph(resource.get("title", "N/A"), styles["Heading2"]))
                story.append(Paragraph(f"ID: {resource_id if resource_id else 'N/A'}", styles["Normal"]))
                story.append(Paragraph(f"Type: {resource.get('type', 'N/A')}", styles["Normal"]))
                story.append(Paragraph(f"Focus Area: {resource.get('focus_area', 'N/A')}", styles["Normal"]))
                story.append(Paragraph(preprocess_html_for_reportlab(resource.get("short_description", "")), styles["Normal"]))

                # Process activities
                if "activities" in resource:
                    # Activities - Heading 3
                    story.append(Paragraph("Activities", styles["Heading3"]))
                    for activity in resource.get("activities", []):
                        # Activity Title - Heading 4
                        story.append(Paragraph(activity.get("title", "N/A"), styles["Heading4"]))
                        story.append(Paragraph(f"ID: {activity.get('id', 'N/A')}", styles["Normal"]))
                        story.append(Paragraph(f"Type: {activity.get('type', 'N/A')}", styles["Normal"]))

                        # Content Blocks
                        if "content_blocks" in activity:
                            # Content Blocks - Heading 5 (if style exists)
                            story.append(Paragraph("Content Blocks", styles.get('Heading5', styles['Heading4'])))
                            for block in activity.get("content_blocks", []):
                                # Block Title - Heading 6 (if style exists)
                                block_title = block.get("title", block.get("type", ""))
                                story.append(Paragraph(block_title, styles.get('Heading6', styles['Heading4'])))
                                story.append(Paragraph(f"Type: {block.get('type', 'N/A')}", styles["Normal"]))
                                if "content" in block:
                                    # ... (existing PDF content block processing) ...
                                    content = block["content"]
                                    if isinstance(content, dict):
                                        if "text" in content:
                                            story.append(Paragraph(preprocess_html_for_reportlab(content['text']), styles["Normal"]))
                                        if "title" in content:
                                            if block_title != content['title']:
                                                story.append(Paragraph(f"Content Title: {preprocess_html_for_reportlab(content['title'])}", styles["Normal"]))
                                        if "columns" in content:
                                            for col in content.get("columns", []):
                                                col_title = preprocess_html_for_reportlab(col.get('title', ''))
                                                col_text = preprocess_html_for_reportlab(col.get('text', ''))
                                                story.append(Paragraph(f"{col_title}: {col_text}", styles["Normal"]))
                                    elif isinstance(content, list):
                                         for item in content:
                                             if isinstance(item, dict) and "text" in item:
                                                 story.append(Paragraph(preprocess_html_for_reportlab(item['text']), styles["Normal"]))

                        # Process standards coding schemes and statements (COMMENTED OUT)
                        # if "standards" in activity:
                        #    story.append(Paragraph("Standards", styles.get('Heading7', styles['Heading4'])))
                        #    for standard in activity.get("standards", []):
                        #        story.append(Paragraph(f"{standard.get('code','N/A')}: {standard.get('statement', 'N/A')}", styles['StandardStyle']))
                                         
                story.append(Spacer(1, 12))
        story.append(PageBreak())

    # Build the PDF
    try:
        doc.build(story)
        print(f"PDF saved as {output_file}")
    except Exception as e:
        print(f"Error building PDF: {e}\nCheck logs for details.", file=sys.stderr)
        raise

# Main function to load JSON and export
def main():
    # Load your JSON data (replace with your file path or string)
    with open("craft_data.json", "r") as f:
        data = json.load(f)
    
    # Export to Word
    export_to_word(data, "preschool_curriculum_llama.docx")
    
    # Export to PDF
    export_to_pdf(data, "preschool_curriculum_llama.pdf")

if __name__ == "__main__":
    main()