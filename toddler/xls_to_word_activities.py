import pandas as pd
from docx import Document
from docx.shared import Pt
from agno.agent import Agent
from agno.models.openai import OpenAIChat
import os

# Initialize OpenAI client (Agno handles this internally via environment variables)
# client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# --- Define Agno Agents ---

language_agent = Agent(
    name="Language Agent",
    role="You are an expert in early childhood language development. Create engaging toddler activities focused on vocabulary, communication, and language skills.",
    model=OpenAIChat(id="gpt-4o"),
)

math_agent = Agent(
    name="Math Agent",
    role="You are an expert in early childhood math education. Create engaging toddler activities focused on counting, shapes, patterns, and basic math concepts.",
    model=OpenAIChat(id="gpt-4o"),
)

science_agent = Agent(
    name="Science Agent",
    role="You are an expert in early childhood science education. Create engaging toddler activities focused on exploration, cause and effect, and natural curiosity.",
    model=OpenAIChat(id="gpt-4o"),
)

art_agent = Agent(
    name="Art Agent",
    role="You are an expert in early childhood art education. Create engaging toddler activities focused on creativity, sensory experiences, and fine motor skills.",
    model=OpenAIChat(id="gpt-4o"),
)

music_agent = Agent(
    name="Music Agent",
    role="You are an expert in early childhood music education. Create engaging toddler activities focused on rhythm, movement, songs, and musical exploration.",
    model=OpenAIChat(id="gpt-4o"),
)

default_agent = Agent(
    name="General Activity Agent",
    role="You are an expert in early childhood education. Create engaging toddler activities focused on development, learning, and social interaction.",
    model=OpenAIChat(id="gpt-4o"),
)

# Create a map for easy agent lookup
agent_map = {
    'Language': language_agent,
    'Math': math_agent,
    'Science': science_agent,
    'Art': art_agent,
    'Music': music_agent,
    'default': default_agent
}

# Load spreadsheet data
def load_spreadsheet(filename):
    data = pd.read_excel(filename, sheet_name=None)
    return data

def clean_activity_text(text):
    """Clean the activity text by removing formatting characters and extra whitespace."""
    if not text:
        return None
    
    try:
        # Remove all box drawing characters and special formatting
        text = text.replace('┃', '').replace('│', '').replace('─', '').replace('━', '')
        text = text.replace('┗', '').replace('┛', '').replace('┏', '').replace('┓', '')
        text = text.replace('┣', '').replace('┫', '').replace('┳', '').replace('┻', '')
        text = text.replace('╋', '').replace('╂', '').replace('╁', '').replace('╀', '')
        
        # Remove any remaining special characters
        text = ''.join(char for char in text if ord(char) < 128)
        
        # Remove extra whitespace and normalize line breaks
        lines = []
        for line in text.split('\n'):
            line = line.strip()
            if line and not line.startswith('---'):  # Skip separator lines
                lines.append(line)
        
        cleaned_text = '\n'.join(lines)
        
        # Debug: Print the cleaned text
        print("\n=== Cleaned Text ===")
        print(cleaned_text)
        print("=== End Cleaned Text ===\n")
        
        return cleaned_text
    except Exception as e:
        print(f"Error cleaning text: {str(e)}")
        return None

def generate_activity_with_agent(agent_map, section_type, activity_summary, theme):
    """Generates an activity using the appropriate Agno agent."""
    
    # Select the correct agent based on the section type
    agent = agent_map.get(section_type, agent_map['default'])
    
    # Construct the prompt for the agent
    prompt = f"""
    Theme: "{theme}"
    Activity Summary: "{activity_summary}"
    
    Generate a detailed activity based on your specialized role. Include:
    1. Clear step-by-step instructions
    2. Materials needed (if any)
    3. Learning objectives
    4. Tips for teachers
    
    Format your response with proper line breaks and clear sections.
    Use markdown-style formatting for sections (e.g., **Objective:**, **Materials Needed:**)
    Do not use any special characters, box drawing characters, or separators.
    """
    
    try:
        # Use print_response method and capture its output
        response = agent.print_response(prompt, stream=False)
        if not response:
            print(f"Warning: Empty response from {section_type} agent")
            return None
        
        # Debug: Print the raw response
        print("\n=== Raw Response ===")
        print(response)
        print("=== End Raw Response ===\n")
        
        # Clean the response text
        cleaned_response = clean_activity_text(response)
        if not cleaned_response:
            print(f"Warning: Empty response after cleaning from {section_type} agent")
            return None
            
        return cleaned_response
    except Exception as e:
        print(f"Error generating activity: {str(e)}")
        return None

# Create a structured Word document
def create_document(activities_dict, output_filename):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Toddler Weekly Activity Plan', 0)
    title.alignment = 1  # Center alignment
    title.style.font.size = Pt(16)
    title.style.font.bold = True

    # Add activities to document
    for theme, sections in activities_dict.items():
        # Theme heading
        theme_heading = doc.add_heading(theme, level=1)
        theme_heading.style.font.size = Pt(14)
        theme_heading.style.font.bold = True
        
        # Add sections
        for section_title, activity_text in sections.items():
            # Section heading
            section_heading = doc.add_heading(section_title, level=2)
            section_heading.style.font.size = Pt(12)
            section_heading.style.font.bold = True
            
            # Activity text
            para = doc.add_paragraph()
            para.style.font.size = Pt(11)
            
            # Format the activity text with proper spacing
            if activity_text:  # Check if activity_text exists
                # Clean the text one more time before adding to document
                cleaned_text = clean_activity_text(activity_text)
                if cleaned_text:
                    # Debug: Print what's being added to the document
                    print(f"\n=== Adding to Document: {section_title} ===")
                    print(cleaned_text)
                    print("=== End Document Addition ===\n")
                    
                    # Process each line with proper formatting
                    for line in cleaned_text.split('\n'):
                        if line.strip():  # Only add non-empty lines
                            # Check for bold sections
                            if line.startswith('**') and line.endswith('**'):
                                run = para.add_run(line.strip('*') + '\n')
                                run.bold = True
                            else:
                                run = para.add_run(line.strip() + '\n')
                            run.font.size = Pt(11)
                else:
                    run = para.add_run(f"Error: Could not process activity text for {section_title}. Please try again.")
                    run.font.size = Pt(11)
            else:
                # Add error message if activity text is missing
                run = para.add_run(f"Error: No activity text generated for {section_title}. Please try again.")
                run.font.size = Pt(11)
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Add page break between themes
        doc.add_page_break()

    # Save the document
    doc.save(output_filename)
    print(f"\nDocument saved to {output_filename}")

# Main logic
def main(spreadsheet_filename, output_doc_filename):
    spreadsheet_data = load_spreadsheet(spreadsheet_filename)
    print("\n=== Loading Spreadsheet ===")
    print("Spreadsheet loaded successfully.")

    # Select first sheet for simplicity (modify as needed)
    sheet_name = list(spreadsheet_data.keys())[0]
    df = spreadsheet_data[sheet_name]
    
    # Get the theme name from the first column
    theme_name = df.iloc[0, 0]
    print(f"\n=== Processing Theme ===")
    print(f"Theme: {theme_name}")
    
    # Create activities dictionary with theme structure
    activities_dict = {theme_name: {}}
    
    # Process each row that contains activity information
    current_section = None
    for idx, row in df.iterrows():
        if pd.isna(row.iloc[0]):
            continue
            
        # Check if this is a section header
        if row.iloc[0] in ['Language', 'Math', 'Science', 'Art', 'Music']:
            current_section = row.iloc[0]
            print(f"\n=== Processing {current_section} Section ===")
            continue
            
        # If we have a current section and activity text, generate the activity
        if current_section and not pd.isna(row.iloc[0]):
            activity_summary = row.iloc[0]
            print(f"\nGenerating activity for {current_section}:")
            print(f"Activity Title: {activity_summary}")
            try:
                activity_text = generate_activity_with_agent(agent_map, current_section, activity_summary, theme_name)
                if activity_text:
                    activities_dict[theme_name][f"{current_section} - {activity_summary}"] = activity_text
                    print("✓ Activity generated successfully")
                    
                    # Save progress after each activity
                    create_document(activities_dict, output_doc_filename)
                    print("✓ Progress saved to document")
                else:
                    print("✗ Failed to generate activity text")
            except Exception as e:
                print(f"Error processing activity: {str(e)}")
                print("Skipping to next activity...")

    # Save final document with all generated activities
    create_document(activities_dict, output_doc_filename)
    # Final confirmation
    print("\n=== Final Document Created ===")
    print(f"All activities have been generated and saved to {output_doc_filename}")

if __name__ == "__main__":
    spreadsheet_filename = 'Toddler _ Themes_for Brandon.xlsx'
    output_doc_filename = 'Generated_Weekly_Activities_agent.docx'

    main(spreadsheet_filename, output_doc_filename)
